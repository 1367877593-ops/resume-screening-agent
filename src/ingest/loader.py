"""pdf / docx / 纯文本 -> RawDoc。

JD 走纯文本路径（页面上是个文本框），简历走文件路径。
解析库用惰性导入：只想跑 demo 的人不必装 pypdf 和 python-docx。
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import List, Optional

from ingest.chunker import chunk_text
from ingest.normalizer import normalize
from schema.document import RawDoc

SUPPORTED = {".pdf", ".docx", ".txt", ".md"}


def _doc_id(seed: str) -> str:
    return hashlib.sha1(seed.encode("utf-8")).hexdigest()[:12]


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError("解析 PDF 需要 pypdf：pip install pypdf") from e
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    try:
        import docx
    except ImportError as e:
        raise RuntimeError("解析 Word 需要 python-docx：pip install python-docx") from e
    d = docx.Document(str(path))
    parts: List[str] = [p.text for p in d.paragraphs]
    # 简历里的关键信息常放在表格中，漏掉表格会丢掉大半内容
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def load_text(text: str, filename: str = "input.txt", doc_id: Optional[str] = None) -> RawDoc:
    full = normalize(text)
    did = doc_id or _doc_id(filename + full[:200])
    return RawDoc(doc_id=did, filename=filename, full_text=full, chunks=chunk_text(did, full))


def content_doc_id(filename: str, data: bytes) -> str:
    """由「原始文件名 + 内容」派生 doc_id，与文件落在哪里无关。

    上传的 pdf/docx 需要落到临时文件才能解析，而临时路径每次都不一样。
    若用路径参与派生，同一份简历每次上传都会得到不同的 doc_id ——
    而 doc_id 会进 prompt，于是缓存永远命不中、模型每次重答、分数跟着漂。
    """
    digest = hashlib.sha1(data).hexdigest()
    return _doc_id(f"{filename}:{digest}")


def load_file(path, filename: Optional[str] = None, doc_id: Optional[str] = None) -> RawDoc:
    """`filename` / `doc_id` 用于上传场景：真实身份由调用方给，不取临时路径。"""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED:
        raise ValueError(f"不支持的文件类型：{suffix}（支持 {sorted(SUPPORTED)}）")

    if suffix == ".pdf":
        raw = _read_pdf(path)
    elif suffix == ".docx":
        raw = _read_docx(path)
    else:
        raw = path.read_text(encoding="utf-8")

    if not raw.strip():
        raise ValueError(
            f"{path.name} 未能提取到任何文本。"
            "扫描件或图片型 PDF 需要 OCR，当前版本不支持。"
        )
    name = filename or path.name
    return load_text(raw, filename=name, doc_id=doc_id or _doc_id(str(path) + raw[:200]))
